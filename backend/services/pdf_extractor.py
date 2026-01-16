import PyPDF2
import docx
from io import BytesIO
import os

def extract_pdf_text(pdf_file) -> str:
    """
    Extract text from PDF file.
    Works with file-like objects from Flask request.files
    """
    try:
        # Reset file pointer if needed
        if hasattr(pdf_file, 'seek'):
            pdf_file.seek(0)
        
        # Read PDF
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        # Extract text from all pages
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        
        return text.strip()
    
    except Exception as e:
        print(f"Error extracting text from PDF: {str(e)}")
        raise

def extract_docx_text(docx_file) -> str:
    """
    Extract text from DOCX file.
    Works with file-like objects from Flask request.files
    """
    try:
        # Reset file pointer if needed
        if hasattr(docx_file, 'seek'):
            docx_file.seek(0)
        
        # Read DOCX
        doc = docx.Document(docx_file)
        text = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        
        return "\n".join(text).strip()
    
    except Exception as e:
        print(f"Error extracting text from DOCX: {str(e)}")
        raise

def extract_text_from_file(file_path_or_fileobj, file_type):
    """
    Extract text from PDF or DOCX file.
    Accepts both file path (str) and file-like object.
    """
    if file_type == 'pdf':
        if isinstance(file_path_or_fileobj, str):
            with open(file_path_or_fileobj, 'rb') as f:
                return extract_pdf_text(f)
        else:
            return extract_pdf_text(file_path_or_fileobj)
    elif file_type == 'docx':
        if isinstance(file_path_or_fileobj, str):
            with open(file_path_or_fileobj, 'rb') as f:
                return extract_docx_text(f)
        else:
            return extract_docx_text(file_path_or_fileobj)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

# LEGACY FUNCTION - REQUIRED BY rag_pipeline.py
def extract_text_from_pdf(pdf_file) -> str:
    """
    Legacy function name for rag_pipeline.py compatibility
    """
    return extract_pdf_text(pdf_file)
